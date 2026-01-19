"""Parser for Microsoft Word documents (.docx and legacy .doc).

This parser prefers the lightweight `python-docx` library for .docx
files and will optionally use `textract` for legacy .doc files when
available.  Extraction is best-effort: paragraphs, table cells and
any other text elements are concatenated into a single document.  If
the necessary dependencies are not installed, the parser logs an
informative error and returns no documents rather than raising.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import List

try:
    import docx  # type: ignore
except ImportError as exc:  # pragma: no cover - optional dependency
    docx = None  # type: ignore

try:
    import textract  # type: ignore
except ImportError as exc:  # pragma: no cover - optional dependency
    textract = None  # type: ignore

from .base_parser import DocumentParser
from ..document import Document


logger = logging.getLogger(__name__)


class WordParser(DocumentParser):
    """Concrete parser for Word documents."""

    def parse(self, file_path: str) -> List[Document]:
        documents: List[Document] = []
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix not in {".docx", ".doc"}:
            logger.error("WordParser: unsupported file extension %s", suffix)
            return documents
        if not path.is_file():
            logger.error("WordParser: %s is not a file", file_path)
            return documents

        if suffix == ".doc":
            return self._parse_doc(path, documents)
        return self._parse_docx(path, documents)

    def _parse_doc(self, path: Path, documents: List[Document]) -> List[Document]:
        if textract is None:
            logger.error(
                "WordParser: textract library is not installed. Please install textract to parse .doc files."
            )
            return documents
        try:
            text_bytes = textract.process(str(path))
            content = text_bytes.decode("utf-8", errors="ignore")
        except Exception:
            logger.exception("WordParser: failed to extract text from .doc file %s", path)
            return documents
        metadata = {"file_path": str(path), "format": "doc"}
        documents.append(Document(text=content, metadata=metadata))
        logger.debug("WordParser: parsed legacy .doc %s", path)
        return documents

    def _parse_docx(self, path: Path, documents: List[Document]) -> List[Document]:
        if docx is None:
            logger.error(
                "WordParser: python-docx library is not installed. Please install python-docx to parse .docx files."
            )
            return documents
        try:
            doc = docx.Document(path)
        except Exception:
            logger.exception("WordParser: failed to open .docx file %s", path)
            return documents

        text_parts: List[str] = []
        for para in doc.paragraphs:
            if para.text:
                text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = "\t".join(filter(None, cells))
                if row_text:
                    text_parts.append(row_text)

        content = "\n".join(text_parts)
        metadata = {"file_path": str(path), "format": "docx"}
        documents.append(Document(text=content, metadata=metadata))
        logger.debug("WordParser: parsed .docx file %s with %d text blocks", path, len(text_parts))
        return documents
